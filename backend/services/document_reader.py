import io
import os
from docx import Document
from pathlib import Path


def read_docx(file_path: str) -> str:
    doc = Document(file_path)
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def read_docx_bytes(data: bytes) -> str:
    doc = Document(io.BytesIO(data))
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
