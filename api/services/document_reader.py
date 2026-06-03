import io
import os
from docx import Document
from pypdf import PdfReader
from pathlib import Path


def read_docx(file_path: str) -> str:
    try:
        doc = Document(file_path)
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    except Exception as e:
        print(f"Error reading docx file: {e}")
        return ""


def read_docx_bytes(data: bytes) -> str:
    try:
        doc = Document(io.BytesIO(data))
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    except Exception as e:
        print(f"Error reading docx bytes: {e}")
        return ""


def read_pdf_bytes(data: bytes) -> str:
    try:
        reader = PdfReader(io.BytesIO(data))
        text = ""
        for page in reader.pages:
            t = page.extract_text()
            if t:
                text += t + "\n"
        return text.strip()
    except Exception as e:
        print(f"Error reading pdf bytes: {e}")
        return ""

